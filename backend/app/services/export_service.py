"""
Dışa aktarma servisi — BibTeX (.bib) ve Word (.docx) üretimi.
LLM gerektirmez; deterministik, hızlı.
"""
from __future__ import annotations

import re
import unicodedata
from io import BytesIO


# ── BibTeX ────────────────────────────────────────────────────────────────────

def _bibtex_key(title: str, author: str | None, year: int | str | None) -> str:
    """smith2024 veya anon2024 formatında anahtar üretir."""
    if author:
        last = author.split(",")[0].split()[-1]
    else:
        last = "anon"

    # ASCII'ye indir, harf dışı karakter at
    last = unicodedata.normalize("NFD", last)
    last = "".join(c for c in last if unicodedata.category(c) != "Mn")
    last = re.sub(r"[^a-zA-Z]", "", last).lower()[:12] or "anon"

    yr = str(year)[:4] if year else "0000"
    return f"{last}{yr}"


def _bibtex_escape(value: str) -> str:
    """BibTeX özel karakterlerini kaçır."""
    return (
        str(value)
        .replace("&",  r"\&")
        .replace("%",  r"\%")
        .replace("$",  r"\$")
        .replace("#",  r"\#")
        .replace("_",  r"\_")
        .replace("{",  r"\{")
        .replace("}",  r"\}")
        .replace("~",  r"\~{}")
        .replace("^",  r"\^{}")
    )


def generate_bibtex(documents: list) -> str:
    """
    Bir belge listesinden BibTeX dosyası içeriği üretir.
    Her belgein citation_data alanını kullanır.
    LLM çağrısı yapmaz.
    """
    entries: list[str] = []
    seen_keys: dict[str, int] = {}

    for doc in documents:
        cd    = doc.citation_data or {}
        title = cd.get("title") or doc.title or "Başlıksız"
        author= cd.get("author") or ""
        year  = cd.get("year")
        pub   = cd.get("publisher") or ""
        doi   = cd.get("doi") or ""

        base_key = _bibtex_key(title, author, year)
        # Çakışma varsa a/b/c ekle
        if base_key in seen_keys:
            seen_keys[base_key] += 1
            key = f"{base_key}{chr(96 + seen_keys[base_key])}"
        else:
            seen_keys[base_key] = 0
            key = base_key

        entry_type = "article" if pub else "misc"

        fields: list[str] = [f"  title     = {{{_bibtex_escape(title)}}}"]
        if author:
            fields.append(f"  author    = {{{_bibtex_escape(author)}}}")
        if year:
            fields.append(f"  year      = {{{year}}}")
        if pub:
            fields.append(f"  journal   = {{{_bibtex_escape(pub)}}}")
        if doi:
            fields.append(f"  doi       = {{{_bibtex_escape(doi)}}}")
        fields.append(f"  note      = {{doc\\_id: {doc.doc_id}}}")

        body = ",\n".join(fields)
        entries.append(f"@{entry_type}{{{key},\n{body}\n}}")

    header = "% ArcticDocs BibTeX Export\n% https://github.com/hayatdiler/ai-document-research-system\n\n"
    return header + "\n\n".join(entries)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_docx(collection_name: str, report_text: str) -> bytes:
    """
    Koleksiyon rapor metninden Word belgesi üretir.
    Markdown benzeri # başlıklar ve - madde işaretleri desteklenir.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Başlık
    title_para = doc.add_heading(f"Koleksiyon Raporu: {collection_name}", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # boşluk

    # Rapor içeriği
    clean = report_text.replace("**", "").replace("*", "")

    for line in clean.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "• ")):
            para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            para.paragraph_format.left_indent = Cm(0.5)
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
